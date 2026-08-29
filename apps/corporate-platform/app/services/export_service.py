"""导出服务 — 生成 docx/pdf/xlsx 文件"""

import asyncio
import os
from datetime import datetime
from io import BytesIO

from loguru import logger


# ---- CJK 字体探测 ----

def _find_cjk_font() -> str:
    """查找系统上可用的 CJK 字体"""
    candidates = [
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        "/usr/share/fonts/truetype/arphic/uming.ttc",
    ]
    for path in candidates:
        if os.path.isfile(path):
            return path
    # 回退：搜索 /usr/share/fonts 下任意 ttf/ttc/otf
    for root, _dirs, files in os.walk("/usr/share/fonts"):
        for f in files:
            if f.endswith((".ttf", ".ttc", ".otf")):
                return os.path.join(root, f)
    return ""


# ---- DOCX ----

async def export_to_docx(content: str, title: str = "") -> bytes:
    """导出纯文本为 DOCX"""
    def _build():
        from docx import Document

        doc = Document()
        doc.add_heading(title or "导出文档", level=1)
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")

        for block in content.split("\n"):
            doc.add_paragraph(block) if block.strip() else doc.add_paragraph("")

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    return await asyncio.to_thread(_build)


# ---- PDF ----

async def export_to_pdf(content: str, title: str = "") -> bytes:
    """导出纯文本为 PDF（支持中文）"""
    font_path = _find_cjk_font()

    def _build():
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # 注册 CJK 字体
        cjk_available = False
        if font_path and os.path.isfile(font_path):
            try:
                pdf.add_font("CJK", "", font_path, uni=True)
                cjk_available = True
            except Exception as e:
                logger.warning(f"CJK 字体注册失败: {e}")

        font_name = "CJK" if cjk_available else "Helvetica"

        # 标题
        pdf.set_font(font_name, size=16)
        pdf.multi_cell(0, 10, title or "导出文档")
        pdf.ln(4)

        # 时间戳
        pdf.set_font(font_name, size=9)
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        pdf.cell(0, 8, f"生成时间: {ts}", ln=True)
        pdf.ln(4)

        # 正文
        pdf.set_font(font_name, size=11)
        for line in content.split("\n"):
            pdf.multi_cell(0, 6, line)

        return pdf.output()

    return await asyncio.to_thread(_build)


# ---- XLSX ----

async def export_to_xlsx(rows: list[dict], title: str = "") -> bytes:
    """导出列表数据为 Excel"""
    def _build():
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill

        wb = Workbook()
        ws = wb.active
        ws.title = (title or "导出数据")[:31]

        if not rows:
            ws.cell(row=1, column=1, value="无数据")
            buf = BytesIO()
            wb.save(buf)
            return buf.getvalue()

        headers = list(rows[0].keys())
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="409EFF", end_color="409EFF", fill_type="solid")
        header_align = Alignment(horizontal="center")

        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align

        for row_idx, row in enumerate(rows, 2):
            for col_idx, header in enumerate(headers, 1):
                ws.cell(row=row_idx, column=col_idx, value=row.get(header, ""))

        # 自适应列宽
        for col_idx, header in enumerate(headers, 1):
            max_len = len(str(header))
            for row in rows:
                val_len = len(str(row.get(header, "")))
                if val_len > max_len:
                    max_len = val_len
            col_letter = ws.cell(row=1, column=col_idx).column_letter
            ws.column_dimensions[col_letter].width = min(max_len + 4, 40)

        buf = BytesIO()
        wb.save(buf)
        return buf.getvalue()

    return await asyncio.to_thread(_build)


# ---- 分发器 ----

CONTENT_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

EXTENSIONS = {"docx": ".docx", "pdf": ".pdf", "xlsx": ".xlsx"}


async def export_document(
    content: str = "",
    fmt: str = "docx",
    title: str = "",
    doc_type: str = "",
    rows: list[dict] | None = None,
) -> tuple[bytes, str, str]:
    """分发器 — 返回 (file_bytes, media_type, filename)"""
    if fmt == "xlsx":
        file_bytes = await export_to_xlsx(rows or [], title)
    elif fmt == "pdf":
        file_bytes = await export_to_pdf(content, title)
    else:
        file_bytes = await export_to_docx(content, title)

    media_type = CONTENT_TYPES.get(fmt, CONTENT_TYPES["docx"])
    ext = EXTENSIONS.get(fmt, ".docx")

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = (title or doc_type or "export")[:30].replace(" ", "_")
    filename = f"{safe_title}_{ts}{ext}"

    return file_bytes, media_type, filename
